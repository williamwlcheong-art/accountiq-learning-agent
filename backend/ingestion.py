"""
PDF / Excel ingestion pipeline:
  1. Extract text (pdfplumber + pytesseract for image pages, pandas for Excel)
  2. Send to OpenAI via structured output with GAAP/IFRS system prompt + pattern hints
  3. Parse structured response → financial_rows + narrative
  4. Record new label patterns for future learning
"""
import os
import json
import asyncio
from typing import Optional

import pdfplumber

try:
    import pytesseract
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False

try:
    from docx import Document as DocxDocument
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

import aiosqlite
from db import record_patterns, get_pattern_library, normalise_label

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5.4-mini")
MAX_TEXT_CHARS = 60_000
OCR_DPI = 300

# Canonical row definitions (P&L + BS)
PNL_ROWS = [
    ("revenue",              "Revenue / Sales"),
    ("cogs",                 "Cost of goods sold (COGS)"),
    ("gross_profit",         "Gross profit"),
    ("wages_salaries",       "Wages, salaries and staff costs"),
    ("rent_occupancy",       "Rent and occupancy costs"),
    ("advertising_marketing", "Advertising and marketing"),
    ("insurance",            "Insurance"),
    ("motor_vehicle_expenses", "Motor vehicle expenses"),
    ("repairs_maintenance",  "Repairs and maintenance"),
    ("admin_professional_fees", "Administration and professional fees"),
    ("other_operating_expenses", "Other operating expenses"),
    ("operating_expenses",   "Total operating expenses"),
    ("ebitda",               "EBITDA"),
    ("depreciation",         "Depreciation & amortisation"),
    ("ebit",                 "EBIT / Operating profit"),
    ("interest_expense",     "Interest / finance costs"),
    ("pbt",                  "Profit before tax"),
    ("tax",                  "Income tax expense"),
    ("net_profit",           "Net profit / NPAT"),
]

BS_ROWS = [
    ("cash_and_bank",           "Cash & bank"),
    ("trade_debtors",           "Trade debtors / receivables"),
    ("inventory",               "Inventory / stock"),
    ("other_current_assets",    "Other current assets"),
    ("total_current_assets",    "Total current assets"),
    ("fixed_assets_net",        "Fixed assets (net PP&E)"),
    ("other_noncurrent_assets", "Other non-current assets"),
    ("total_assets",            "Total assets"),
    ("trade_creditors",         "Trade creditors / payables"),
    ("short_term_debt",         "Short-term debt / current borrowings"),
    ("other_current_liab",      "Other current liabilities"),
    ("total_current_liab",      "Total current liabilities"),
    ("long_term_debt",          "Long-term debt / non-current borrowings"),
    ("other_noncurrent_liab",   "Other non-current liabilities"),
    ("total_liabilities",       "Total liabilities"),
    ("shareholders_equity",     "Shareholders equity / net assets"),
]

CF_ROWS = [
    ("operating_cashflow",  "Cash flows from operating activities"),
    ("investing_cashflow",  "Cash flows from investing activities"),
    ("financing_cashflow",  "Cash flows from financing activities"),
    ("net_change_in_cash",  "Net change in cash and cash equivalents"),
]

EQ_ROWS = [
    ("opening_equity",          "Opening equity / balance at beginning"),
    ("net_profit",              "Net profit for the period"),
    ("dividends_paid",          "Dividends / distributions paid"),
    ("other_equity_movements",  "Other equity movements"),
    ("closing_equity",          "Closing equity / balance at end"),
]

ALL_ROWS = (
    [("pnl", k, lbl) for k, lbl in PNL_ROWS]
    + [("bs",  k, lbl) for k, lbl in BS_ROWS]
    + [("cf",  k, lbl) for k, lbl in CF_ROWS]
    + [("eq",  k, lbl) for k, lbl in EQ_ROWS]
)

# ---------------------------------------------------------------------------
# System prompt — GAAP / IFRS methodology
# ---------------------------------------------------------------------------

SYSTEM_PROMPT = """You are a financial data extraction specialist with deep expertise in both GAAP (US Generally Accepted Accounting Principles) and IFRS (International Financial Reporting Standards).

## Detecting the Reporting Standard
Look for these signals:

GAAP indicators:
- SEC filings (10-K, 10-Q, 8-K), US-based registrant
- "Common stock", "Additional paid-in capital", "Selling, general and administrative (SG&A)"
- Auditor references to "US GAAP"

IFRS indicators:
- References to IAS/IFRS standard numbers (e.g. IFRS 15, IAS 36)
- "Share capital", "Share premium", "Turnover", "Profit for the year"
- Non-US exchange listings (NZX, ASX, LSE, SGX, etc.)
- Auditor opinion referencing "true and fair view"

## GAAP vs IFRS Terminology (normalise to canonical keys regardless of label used)
| Canonical key       | GAAP labels                        | IFRS labels                              |
|---------------------|------------------------------------|------------------------------------------|
| revenue             | Net revenues, Net sales            | Revenue, Turnover, Net revenue           |
| cogs                | Cost of goods sold                 | Cost of sales, Cost of revenue           |
| gross_profit        | Gross profit                       | Gross profit                             |
| ebit                | Operating income, Income from ops  | Operating profit, Profit from operations |
| net_profit          | Net income, Net earnings           | Profit for the year, Net profit, NPAT    |
| shareholders_equity | Total stockholders' equity         | Total equity, Net assets                 |

## Canonical P&L Keys (statement: "pnl")
revenue, cogs, gross_profit, wages_salaries, rent_occupancy, advertising_marketing,
insurance, motor_vehicle_expenses, repairs_maintenance, admin_professional_fees,
other_operating_expenses, operating_expenses, ebitda, depreciation, ebit,
interest_expense, pbt, tax, net_profit

## Canonical Balance Sheet Keys (statement: "bs")
cash_and_bank, trade_debtors, inventory, other_current_assets, total_current_assets,
fixed_assets_net, other_noncurrent_assets, total_assets, trade_creditors, short_term_debt,
other_current_liab, total_current_liab, long_term_debt, other_noncurrent_liab,
total_liabilities, shareholders_equity

## Canonical Cash Flow Keys (statement: "cf")
operating_cashflow, investing_cashflow, financing_cashflow, net_change_in_cash

## Canonical Equity Changes Keys (statement: "eq")
opening_equity, net_profit, dividends_paid, other_equity_movements, closing_equity

## Extraction Rules
1. Preserve the exact original label in raw_label — this is used for pattern learning
2. Values as plain numbers — no commas, no currency symbols
3. Negative values use negative numbers; dashes or blanks = null
4. If stated in thousands/millions, set unit accordingly and keep values in that unit
5. NEVER fabricate values — use null if not found
6. Assign confidence 0–1 based on how clearly the value maps to the canonical key
7. SIGN CONVENTION: Cost/expense keys (cogs, wages_salaries, rent_occupancy,
   advertising_marketing, insurance, motor_vehicle_expenses, repairs_maintenance,
   admin_professional_fees, other_operating_expenses, operating_expenses, depreciation,
   interest_expense, tax) must be returned as NEGATIVE numbers.
   Revenue and asset/equity keys must be POSITIVE.
   A post-processing normalisation layer enforces this, but supply the correct sign.
8. BALANCE-SHEET CLASSIFICATION: Map only actual plant, equipment, vehicles or other
   identified tangible operating assets to fixed_assets_net. Do not map a total
   non-current-assets line to fixed assets; use other_noncurrent_assets unless the
   source breaks out the fixed-asset amount. Map only stated loans, borrowings,
   mortgages, hire purchase, leases or shareholder loans to debt. Do not map a
   total non-current-liabilities line to long_term_debt; use other_noncurrent_liab
   unless the source identifies the debt component.

## Narrative
Write a 3–4 paragraph executive summary covering: revenue performance,
profitability (margins), balance sheet position, key highlights and risks.
Use plain business language. Be concise and specific."""


# ---------------------------------------------------------------------------
# Structured-output schema — financial extraction
# ---------------------------------------------------------------------------

_ROW_SCHEMA = {
    "type": "object",
    "required": ["statement", "canonical_key", "raw_label", "values", "confidence"],
    "properties": {
        "statement":     {"type": "string", "enum": ["pnl", "bs", "cf", "eq"]},
        "canonical_key": {"type": "string", "description": "One of the canonical keys listed in the system prompt"},
        "raw_label":     {"type": "string", "description": "Exact label as it appears in the document"},
        "values":        {"type": "object",  "description": "Period → value map, e.g. {\"2025\": 1234567, \"2024\": null}"},
        "confidence":    {"type": "number",  "description": "0–1 confidence score for this mapping"},
    },
}

EXTRACT_RESPONSE_SCHEMA = {
    "type": "object",
    "required": ["periods", "currency", "unit", "reporting_standard", "rows", "narrative"],
    "properties": {
        "periods":            {"type": "array", "items": {"type": "string"}, "description": "Fiscal years found, most recent first"},
        "currency":           {"type": "string", "description": "ISO 4217 code e.g. NZD, AUD, USD"},
        "unit":               {"type": "string", "enum": ["whole", "thousands", "millions"]},
        "reporting_standard": {"type": "string", "enum": ["GAAP", "IFRS", "UNKNOWN"]},
        "rows":               {"type": "array", "items": _ROW_SCHEMA},
        "extraction_notes":   {"type": "string"},
        "narrative":          {"type": "string", "description": "3–4 paragraph executive summary"},
    },
}


# ---------------------------------------------------------------------------
# Text extraction — PDF
# ---------------------------------------------------------------------------

def _page_has_text(page) -> bool:
    text = page.extract_text() or ""
    if len(text.strip()) <= 100:
        return False

    # Some accounting PDFs expose glyph IDs rather than readable text. Treat
    # repeated CID markers as an extraction failure so the page goes through OCR.
    if text.count("(cid:") >= 3:
        return False

    return True


def _ocr_page(page) -> str:
    if not HAS_TESSERACT:
        return ""
    img = page.to_image(resolution=OCR_DPI).original
    return pytesseract.image_to_string(img, config="--psm 6")


def extract_pdf_text(filepath: str) -> tuple[str, list[str], int, bool]:
    """Returns (extraction_text, all_page_texts, page_count, used_ocr)."""
    all_pages: list[str] = []
    used_ocr = False

    with pdfplumber.open(filepath) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            if _page_has_text(page):
                t = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
            else:
                t = _ocr_page(page)
                if t.strip():
                    used_ocr = True
            all_pages.append(t)

    # Score pages — include all with score > 0, in document order (D-05/D-06/D-07)
    try:
        from rule_extractor import _score_page, PNL_SYNS, BS_SYNS, CF_SYNS, EQ_SYNS
        scored = [
            (
                _score_page(pt, PNL_SYNS) + _score_page(pt, BS_SYNS)
                + _score_page(pt, CF_SYNS) + _score_page(pt, EQ_SYNS),
                i, pt
            )
            for i, pt in enumerate(all_pages)
        ]
        # D-05: filter score > 0 (exclude cover pages, directors reports), sort by page index
        selected = [(s, i, pt) for s, i, pt in scored if s > 0]
        selected.sort(key=lambda x: x[1])
        # D-06: if total chars exceed 60K cap, drop lowest-scored pages first
        total_chars = sum(len(f"--- PAGE {i+1} ---\n{pt}") for _, i, pt in selected)
        while total_chars > MAX_TEXT_CHARS and len(selected) > 1:
            min_idx = min(range(len(selected)), key=lambda k: selected[k][0])
            removed = selected.pop(min_idx)
            total_chars -= len(f"--- PAGE {removed[1]+1} ---\n{removed[2]}")
        extraction_parts = [(i, f"--- PAGE {i+1} ---\n{pt}") for _, i, pt in selected]
        extraction_parts.sort(key=lambda x: x[0])
        extraction_text = "\n".join(c for _, c in extraction_parts)
    except ImportError:
        extraction_text = ""

    if not extraction_text:
        extraction_text = "\n".join(f"--- PAGE {i+1} ---\n{p}" for i, p in enumerate(all_pages))
        extraction_text = extraction_text[:MAX_TEXT_CHARS]

    return extraction_text, all_pages, page_count, used_ocr


# ---------------------------------------------------------------------------
# Text extraction — Excel
# ---------------------------------------------------------------------------

def extract_excel_text(filepath: str) -> tuple[str, list[str], int, bool]:
    """Returns (extraction_text, sheet_texts, sheet_count, used_ocr=False)."""
    try:
        import pandas as pd
    except ImportError:
        raise ImportError("pandas is required for Excel ingestion: pip install pandas openpyxl")

    xl = pd.ExcelFile(filepath)
    sheets: list[str] = []
    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name, header=None)
        if df.empty:
            continue
        sheets.append(f"--- SHEET: {sheet_name} ---\n{df.to_string(index=False)}")

    combined = "\n\n".join(sheets)
    return combined[:MAX_TEXT_CHARS], sheets, len(xl.sheet_names), False


# ---------------------------------------------------------------------------
# Text extraction — Word (.docx)
# ---------------------------------------------------------------------------

def extract_docx_text(filepath: str) -> tuple[str, list[str], int, bool]:
    """Returns (extraction_text, [combined], 1, used_ocr=False).
    Table-first extraction: tab-separated cells per row, prefixed with TABLE markers.
    Non-table paragraphs appended after all tables. Merged cells deduplicated by _tc identity.
    """
    if not HAS_PYTHON_DOCX:
        raise ImportError("python-docx is required for Word ingestion: pip install python-docx")

    try:
        doc = DocxDocument(filepath)
        parts: list[str] = []

        for i, table in enumerate(doc.tables):
            parts.append(f"--- TABLE {i+1} ---")
            for row in table.rows:
                # Deduplicate horizontally merged cells by underlying XML element identity.
                # row.cells always returns grid_cols objects — merged cells repeat the same
                # object for each spanned column. Without dedup: "2025\t2025" instead of "2025".
                seen: dict[int, bool] = {}
                cells_text: list[str] = []
                for cell in row.cells:
                    cell_id = id(cell._tc)
                    if cell_id not in seen:
                        seen[cell_id] = True
                        cells_text.append(cell.text.strip())
                parts.append("\t".join(cells_text))

        # Append non-table paragraphs after all tables (D-12)
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)

        combined = "\n".join(parts)
        return combined[:MAX_TEXT_CHARS], [combined], 1, False
    except Exception as e:
        raise RuntimeError(f"Failed to parse Word document: {e}") from e


# ---------------------------------------------------------------------------
# Pattern hints builder
# ---------------------------------------------------------------------------

def _build_pattern_hints(pattern_lib: dict) -> str:
    hints = []
    for stmt in ("pnl", "bs"):
        for key, labels in (pattern_lib.get(stmt) or {}).items():
            top = labels[:6]
            hints.append(f"  {key} ({stmt}): {', '.join(top)}")
    return "\n".join(hints) if hints else "  (none yet — this is the first ingestion)"


# ---------------------------------------------------------------------------
# OpenAI extraction — structured output
# ---------------------------------------------------------------------------

async def call_openai(
    pdf_text: str,
    pattern_lib: dict,
    entity_type: str,
    fiscal_year_end: str,
    model: str = None,
) -> dict:
    """Call OpenAI Responses API and return the validated JSON-shaped output."""
    key = os.environ.get("OPENAI_API_KEY") or OPENAI_API_KEY
    model = model or os.environ.get("OPENAI_MODEL") or OPENAI_MODEL
    if not key:
        raise RuntimeError("OPENAI_API_KEY not set — cannot extract financial statements")
    try:
        from openai import OpenAI
    except ImportError as exc:
        raise RuntimeError(
            "The OpenAI SDK is not installed. Install backend requirements before using live AI extraction."
        ) from exc
    client = OpenAI(api_key=key)

    pattern_hints = _build_pattern_hints(pattern_lib)
    entity_desc = "listed company annual report" if entity_type == "listed" else "SME / private compilation report"

    user_message = f"""Extract all financial data from this {entity_desc}.

Fiscal year end: {fiscal_year_end or 'unknown'}

Learned label patterns (use as hints for mapping raw labels to canonical keys):
{pattern_hints}

Financial statement text:
{pdf_text}"""

    loop = asyncio.get_running_loop()
    response = await loop.run_in_executor(None, lambda: client.responses.create(
        model=model,
        instructions=SYSTEM_PROMPT,
        input=user_message,
        max_output_tokens=4096,
        text={
            "format": {
                "type": "json_schema",
                "name": "financial_extraction",
                "schema": EXTRACT_RESPONSE_SCHEMA,
                "strict": False,
            }
        },
    ))

    raw_text = str(getattr(response, "output_text", "") or "").strip()
    if not raw_text:
        raise RuntimeError("OpenAI did not return a financial extraction")
    try:
        parsed = json.loads(raw_text)
    except json.JSONDecodeError as exc:
        raise RuntimeError("OpenAI returned invalid financial extraction JSON") from exc
    if not isinstance(parsed, dict):
        raise RuntimeError("OpenAI financial extraction must be a JSON object")
    return parsed


# ---------------------------------------------------------------------------
# Sign normalisation
# ---------------------------------------------------------------------------

_COST_KEYS = frozenset({"cogs", "operating_expenses", "depreciation", "interest_expense", "tax"})


def _normalize_signs(rows: list[dict]) -> list[dict]:
    """Return new row dicts with sign-corrected values for known cost keys.
    Only flips strictly positive values; leaves zero and None unchanged.
    Pure function — does not modify rows in place.
    """
    result = []
    for row in rows:
        key = row.get("canonical_key", "")
        if key in _COST_KEYS:
            new_values = {}
            for period, val in row.get("values", {}).items():
                if val is not None and val > 0:
                    new_values[period] = -val
                else:
                    new_values[period] = val
            result.append({**row, "values": new_values})
        else:
            result.append(row)
    return result


# ---------------------------------------------------------------------------
# Database persistence
# ---------------------------------------------------------------------------

async def persist_extraction(
    db: aiosqlite.Connection,
    document_id: int,
    company_id: int,
    parsed: dict,
    entity_type: str,
    exchange: Optional[str],
):
    periods  = parsed.get("periods", [])
    currency = parsed.get("currency", "NZD")
    unit     = parsed.get("unit", "whole")
    rows     = _normalize_signs(parsed.get("rows", []))

    new_patterns = []

    for row in rows:
        stmt    = row.get("statement", "")
        key     = row.get("canonical_key", "")
        raw_lbl = row.get("raw_label", "")
        values  = row.get("values", {})
        conf    = row.get("confidence", 0.8)

        display_lbl = next((lbl for _, k, lbl in ALL_ROWS if k == key), key)

        for period, value in values.items():
            await db.execute("""
                INSERT INTO financial_rows
                    (document_id, company_id, statement, row_key, row_label,
                     period, value, currency, unit, source_text, confidence)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (document_id, company_id, stmt, key, display_lbl,
                  period, value, currency, unit, raw_lbl, conf))

        if raw_lbl and key:
            new_patterns.append({
                "canonical_key": key,
                "statement":     stmt,
                "raw_label":     raw_lbl,
                "entity_type":   entity_type,
                "exchange":      exchange,
            })

    # Save narrative and reporting standard
    narrative          = parsed.get("narrative", "")
    reporting_standard = parsed.get("reporting_standard", "UNKNOWN")
    await db.execute("""
        UPDATE documents SET
            narrative          = ?,
            reporting_standard = ?,
            updated_at         = datetime('now')
        WHERE id = ?
    """, (narrative, reporting_standard, document_id))

    await db.commit()
    await record_patterns(db, new_patterns)

    return len(rows)


# ---------------------------------------------------------------------------
# Main ingestion entry point
# ---------------------------------------------------------------------------

async def ingest_document(
    db: aiosqlite.Connection,
    document_id: int,
    company_id: int,
    filepath: str,
    entity_type: str,
    exchange: Optional[str],
    fiscal_year_end: str,
) -> dict:
    async def log(level: str, msg: str):
        await db.execute(
            "INSERT INTO extraction_log (document_id, level, message) VALUES (?,?,?)",
            (document_id, level, msg)
        )
        await db.commit()
        print(f"[{level.upper()}] doc={document_id}: {msg}")

    result = {"document_id": document_id, "rows_saved": 0, "error": None}

    try:
        await db.execute(
            "UPDATE documents SET extraction_status='processing', updated_at=datetime('now') WHERE id=?",
            (document_id,)
        )
        await db.commit()

        # 1. Extract text (PDF or Excel)
        await log("info", f"Extracting text from {filepath}")
        fp_lower = filepath.lower()
        loop = asyncio.get_running_loop()
        if fp_lower.endswith((".xlsx", ".xls", ".xlsm")):
            extraction_text, all_page_texts, page_count, used_ocr = await loop.run_in_executor(
                None, extract_excel_text, filepath
            )
        elif fp_lower.endswith(".docx"):
            extraction_text, all_page_texts, page_count, used_ocr = await loop.run_in_executor(
                None, extract_docx_text, filepath
            )
        else:
            extraction_text, all_page_texts, page_count, used_ocr = await loop.run_in_executor(
                None, extract_pdf_text, filepath
            )

        await db.execute(
            "UPDATE documents SET page_count=?, has_ocr=?, updated_at=datetime('now') WHERE id=?",
            (page_count, int(used_ocr), document_id)
        )
        await db.commit()
        await log("info", f"Extracted {page_count} pages/sheets ({len(extraction_text)} chars)")

        # 2. Load pattern library
        pattern_lib = await get_pattern_library(db)
        await log("info", f"Pattern library: {sum(len(v) for v in pattern_lib.values())} total patterns")

        # 3. Try OpenAI structured extraction; fall back to rule-based on auth/billing errors
        parsed = None
        extraction_method = "rule_based"
        live_model = os.environ.get("OPENAI_MODEL") or OPENAI_MODEL
        api_key = os.environ.get("OPENAI_API_KEY") or OPENAI_API_KEY

        if api_key and not api_key.startswith("sk-YOUR"):
            try:
                await log("info", f"Calling OpenAI ({live_model}) with structured extraction…")
                parsed = await call_openai(
                    extraction_text, pattern_lib, entity_type, fiscal_year_end, live_model
                )
                extraction_method = live_model
                periods = parsed.get("periods", [])
                std     = parsed.get("reporting_standard", "UNKNOWN")
                await log("info", f"OpenAI extracted {len(parsed.get('rows',[]))} rows for {periods} [{std}]")

                await db.execute("""
                    UPDATE documents SET
                        raw_provider_response = ?,
                        extraction_model    = ?,
                        updated_at          = datetime('now')
                    WHERE id = ?
                """, (json.dumps(parsed), live_model, document_id))
                await db.commit()

            except Exception as openai_err:
                err_str = str(openai_err)
                if any(x in err_str.lower() for x in ("credit balance", "incorrect api key", "authentication", "insufficient_quota")):
                    await log("warn", f"OpenAI unavailable ({err_str[:120]}) — falling back to rule-based extractor")
                    parsed = None
                else:
                    raise

        # Rule-based fallback
        if parsed is None:
            from rule_extractor import rule_based_extract
            parsed = rule_based_extract(all_page_texts)
            extraction_method = "rule_based"
            periods = parsed.get("periods", [])
            await log("info", f"Rule-based extractor: {len(parsed.get('rows',[]))} rows for {periods}")
            await db.execute(
                "UPDATE documents SET extraction_model=?, updated_at=datetime('now') WHERE id=?",
                (extraction_method, document_id)
            )
            await db.commit()

        # 4. Persist rows, patterns, narrative
        rows_saved = await persist_extraction(
            db, document_id, company_id, parsed, entity_type, exchange
        )
        result["rows_saved"] = rows_saved
        result["periods"]    = parsed.get("periods", [])

        # 5. Mark done
        avg_conf = sum(r.get("confidence", 0.8) for r in parsed.get("rows", [])) / max(len(parsed.get("rows", [])), 1)
        await db.execute("""
            UPDATE documents SET
                extraction_status = 'done',
                confidence_score  = ?,
                updated_at        = datetime('now')
            WHERE id = ?
        """, (avg_conf, document_id))
        await db.commit()
        await log("info", f"Done ({extraction_method}). {rows_saved} rows, avg conf {avg_conf:.2f}")

    except Exception as e:
        result["error"] = str(e)
        await db.execute(
            "UPDATE documents SET extraction_status='failed', updated_at=datetime('now') WHERE id=?",
            (document_id,)
        )
        await db.commit()
        await log("error", str(e))
        raise

    return result
